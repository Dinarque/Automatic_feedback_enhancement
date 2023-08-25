# -*- coding: utf-8 -*-
"""
Created on Sun Aug 20 12:32:29 2023

@author: 3b13j
"""
import streamlit as st
import tqdm
from stqdm import stqdm
from gpt_pipeline import label_analysis, update_cost, write_synthesis
import os 
from logs import create_folder
from fpdf import FPDF
from gpt_pipeline import wordreference_link, correction_prompt

def get_data(session_state) :
    
    for k in session_state.chunk_analysis.keys() : 
        if type (session_state.chunk_analysis[k]) is not dict :
                 try : 
                     val = session_state.chunk_analysis[k].replace ("'s ", " s ")
                     val = val.replace ("l' ", "l ")
                     val = val.replace("l'", " l")
                     session_state.chunk_analysis.pop(k)
                     k = k.replace ("l' ", "l ")
                     k =  k.replace ("'s ", " s ")
                    
                     
                     session_state.chunk_analysis[k] = eval(session_state.chunk_analysis[k])
                 except : 
                    answer, cb = correction_prompt(st.session_state.chunks[k])
                    update_cost(session_state, cb)
                    session_state.chunk_analysis[k] = answer
                    st.experimental_rerun()
    ## analyse  
    idx = [int(k) for k in session_state.chunk_analysis ]
    tp = [type (el) for el in idx ]
    tp2 =  tp = [type (el) for el in session_state.chunk_analysis.keys() ]
    
  
    idx_gr = [ k for k in idx if session_state.chunk_analysis[k]["Type"] == "grammar"] 
    idx_voc = [el for el in idx if el not in idx_gr]
    # extract grammar points
    
    labels = [session_state.chunk_analysis[k]["Label"] for k in idx_gr]
    
    
    
    ## display
    st.subheader( " Here are some data about the mistakes")
    st.markdown("Index of grammar mistakes")
    st.markdown(idx_gr)
    st.markdown("Index of Vocabulary mistales")
    st.markdown(idx_voc)
    
    st.write("___")
    st.markdown("Analysis of the grammar themes evoked : ")
    if st.button("Launch label_analysis") : 
        
        
        if "labels" not in session_state or session_state.labels == None : 
            dic_labels , cb = label_analysis(labels)
            update_cost(session_state, cb)
            session_state.labels = dic_labels
        dic_labels = eval(session_state.labels)
        st.subheader("Classification of the student's grammar mistakes")
        st.write(dic_labels)
        st.write("___")
        lab_to_idx = {}
        for k in dic_labels.keys() :
            lab_to_idx[k] = []
        for i in idx_gr :
            for k in dic_labels.keys() :
                if st.session_state.chunk_analysis[i]["Label"] in dic_labels[k] :
                    lab_to_idx[k].append(i)
        
        
        
        lab_to_sent = {} 
        for k in lab_to_idx.keys():
            lab_to_sent[k] = [] 
            for i in  lab_to_idx[k] :
                lab_to_sent[k].append(session_state.chunks[i].sentence)
        st.subheader("Link to the sentences")
        st.write (lab_to_sent)
        st.write("___")
        
        lab_to_lesson = {}
    
        
            
   # if st.button("Create lessons ") :          CREATE STQDME 
        
        if "labels" in session_state and "synthesis" not in session_state or session_state.synthesis == {}:
            
                topics = list(dic_labels.keys() )
                st.subheader ("proceed to synthesis")
                st.write(topics)
                
                for top in topics : st.write(top)
                
                
                for i in stqdm(range(len(topics))):
                    top = topics[i]
                    synth, cb = write_synthesis(top, session_state)
                    update_cost(session_state, cb)
                    lab_to_lesson[top] = synth
                session_state.synthesis = lab_to_lesson
                
        st.subheader ("Here is synthesis")
        
        st.write(session_state.synthesis)
        st.write("___")
        
    # extract vocab
    st.subheader("Analysis of the vocabulary")
    words = {}
    words["en"] = []
    words["fr"] = []
    for i in idx_voc :
        d = session_state.chunk_analysis[i]["Label"]
        words["en"].append(d["en"])
        words["fr"].append(d["fr"])
    words["en"].sort()
    words["fr"].sort()
    session_state.vocabulary = words
    st.write(words)
    st.write("___")

"""

def create_review(session_state , lab_to_lesson ,file_name):
    from docx import Document
    from docx.shared import Inches
    
    
      
    doc = Document() 
    doc.add_heading("Review document" , 0)
    p = doc.add_paragraph(session_state.summary)
    
    doc.add_page_break()
    
    doc.add_heading("Grammar : ", level = 1)
    
    for k in lab_to_lesson.keys():
        doc.add_heading(k+" :", level =2)
        pk = doc.add_paragraph(lab_to_lesson[k])
    
    doc.save('review.docx')

"""
def create_review(session_state , file_name= 'Review document'):
    
    import docx
    doc = docx.Document()
    doc.add_heading(file_name, 0)
    p = doc.add_paragraph(session_state.summary)
    doc.paragraphs[0].runs[0].add_break()
    doc.add_heading('I) Study of all the individual mistakes',0)
    
    for i in range(len(session_state.chunks)) :
        doc.add_heading(f'Mistake {i}', 1)
        doc.add_paragraph(session_state.chunks[i].sentence)
        if i in session_state.chunk_analysis.keys():
            for k in session_state.chunk_analysis[i].keys() :
                doc.add_heading( str(k)+" :",2)
                doc.add_paragraph(str(session_state.chunk_analysis[i][k]))
                
    if "synthesis" in session_state :
    
        doc.add_heading('I) Recap on the grammar points to study',0)
        for lab in session_state.synthesis.keys() :
            doc.add_heading(f'{lab} :',1)
            doc.add_paragraph(session_state.synthesis[lab])
            
    if "vocabulary" in session_state :
    
        doc.add_heading('III) Vocabulary to study',0)
        
        voc = session_state.vocabulary["fr"]
        for el in voc : 
            for e in el : 
                doc.add_paragraph(f'{e} : {wordreference_link(e, "fr")}')
        
            
            
                
    


    doc.save('review_document.docx')
    

    


















